#!/usr/bin/env python3
"""Render scholar authorship audit Markdown and DOCX reports.

Input is a UTF-8 JSON file with this shape:

{
  "report_title": "Report title",
  "executive_summary": ["paragraph 1", "paragraph 2"],
  "identity_and_anchors": ["paragraph 1"],
  "timeline": [{"period": "2010-2014", "stage": "B.S., Example University"}],
  "articles": [{"year": 2024, "title": "...", "journal": "Nature", "role": "corresponding author"}],
  "output_md": "path/to/report.md",
  "output_docx": "path/to/report.docx",
  "language": "en"
}
"""

from __future__ import annotations

import argparse
import copy
import re
import json
import shutil
import sys
from datetime import date
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - environment error path
    raise SystemExit("python-docx is required to render DOCX reports") from exc


SKILL_DIR = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "authorship-audit-template.docx"
DEFAULT_OUTPUT_ROOT = Path.cwd() / "outputs" / "scholar-authorship-auditor"

INVALID_JOURNAL_VALUES = {
    "",
    "journal pending verification",
    "venue pending verification",
    "source pending verification",
    "journal information pending verification",
    "venue information pending verification",
    "\u671f\u520a\u4fe1\u606f\u5f85\u6838\u9a8c",
    "\u671f\u520a\u5f85\u6838\u9a8c",
    "\u6765\u6e90\u5f85\u6838\u9a8c",
    "\u51fa\u7248\u6e90\u5f85\u6838\u9a8c",
}

NON_FINAL_PUBLICATION_MARKERS = {
    "preprint",
    "posted-content",
    "posted_content",
    "discussion",
    "discussion-paper",
    "discussion_paper",
    "supplement",
    "withdrawn",
    "submitted",
    "under-review",
    "under_review",
}

NON_FINAL_JOURNAL_PATTERNS = (
    " discuss",
    "discuss.",
    "discussion paper",
    "preprint",
    "posted content",
    "supplement",
    "withdrawn",
    "not accepted",
)


LABELS = {
    "en": {
        "executive": "Executive summary",
        "identity": "Scholar identity and anchors",
        "timeline": "Education and career timeline",
        "articles": "Qualifying first/co-first/corresponding/co-corresponding journal articles",
        "period": "Period",
        "stage": "Stage",
        "year": "Year",
        "title": "Title",
        "journal": "Journal",
        "role": "Scholar role",
    },
    "zh": {
        "executive": "\u7ed3\u679c\u6458\u8981",
        "identity": "\u8eab\u4efd\u951a\u70b9",
        "timeline": "\u6559\u80b2\u4e0e\u804c\u4e1a\u5c65\u5386",
        "articles": "\u7b2c\u4e00\u4f5c\u8005/\u5171\u540c\u7b2c\u4e00\u4f5c\u8005/\u901a\u8baf\u4f5c\u8005/\u5171\u540c\u901a\u8baf\u4f5c\u8005\u8bba\u6587",
        "period": "\u9636\u6bb5",
        "stage": "\u5c65\u5386",
        "year": "\u5e74\u4efd",
        "title": "\u8bba\u6587\u9898\u540d",
        "journal": "\u671f\u520a",
        "role": "\u5b66\u8005\u89d2\u8272",
    },
}


def as_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    return [text] if text else []


def normalize_language(value: Any) -> str:
    text = str(value or "en").lower()
    return "zh" if text.startswith("zh") or "chinese" in text else "en"


def require_payload(payload: dict[str, Any], key: str) -> Any:
    if key not in payload:
        raise SystemExit(f"Missing required JSON field: {key}")
    return payload[key]


def slugify(value: Any) -> str:
    text = str(value or "").strip().lower()
    if not text:
        return "scholar"
    text = re.sub(r"\s+", "-", text)
    text = re.sub(r"[^a-z0-9._-]+", "-", text)
    text = re.sub(r"-{2,}", "-", text).strip("-._")
    return text or "scholar"


def default_output_paths(payload: dict[str, Any]) -> tuple[Path, Path]:
    scholar_name = (
        payload.get("scholar_name")
        or payload.get("target_name")
        or payload.get("canonical_name")
        or payload.get("report_title")
        or "scholar"
    )
    slug = slugify(scholar_name)
    run_dir = DEFAULT_OUTPUT_ROOT / f"{date.today():%Y%m%d}-{slug}"
    return (
        run_dir / f"{slug}-authorship-audit-report.md",
        run_dir / f"{slug}-authorship-audit-report.docx",
    )


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def row_value(row: Any, *keys: str) -> str:
    if isinstance(row, dict):
        for key in keys:
            value = row.get(key)
            if value is not None and str(value).strip():
                return str(value).strip()
        return ""
    if isinstance(row, (list, tuple)):
        values = [str(item).strip() for item in row]
        return values[0] if values else ""
    return str(row).strip()


def normalized_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


def publication_marker(value: Any) -> str:
    return normalized_text(value).replace(" ", "_")


def is_invalid_journal(value: Any) -> bool:
    journal = normalized_text(value)
    if journal in INVALID_JOURNAL_VALUES:
        return True
    return any(pattern in f" {journal}" for pattern in NON_FINAL_JOURNAL_PATTERNS)


def is_non_final_publication(row: dict[str, Any]) -> bool:
    publication_type = publication_marker(row.get("publication_type") or row.get("type"))
    publication_status = publication_marker(row.get("publication_status") or row.get("status"))
    markers = {publication_type, publication_status}
    return any(marker in NON_FINAL_PUBLICATION_MARKERS for marker in markers if marker)


def validate_articles(payload: dict[str, Any]) -> None:
    errors: list[str] = []
    for idx, row in enumerate(payload.get("articles", []), start=1):
        if isinstance(row, dict):
            title = row_value(row, "title", "paper_title") or f"article #{idx}"
            journal = row_value(row, "journal", "venue", "source")
        elif isinstance(row, (list, tuple)):
            values = [str(item).strip() for item in row]
            title = values[1] if len(values) > 1 and values[1] else f"article #{idx}"
            journal = values[2] if len(values) > 2 else ""
        else:
            errors.append(f"article #{idx}: unsupported article row format")
            continue
        if is_invalid_journal(journal):
            errors.append(
                f"{title}: final report journal/venue is unresolved or non-final ({journal!r})"
            )
        if isinstance(row, dict) and is_non_final_publication(row):
            errors.append(
                f"{title}: publication_type/status is non-final; merge to a published version or exclude it"
            )
    if errors:
        joined = "\n- ".join(errors)
        raise SystemExit(
            "Invalid final article payload. Resolve publication status and venue before rendering:\n"
            f"- {joined}"
        )


def timeline_rows(payload: dict[str, Any]) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    for row in payload.get("timeline", []):
        if isinstance(row, dict):
            period = row_value(row, "period", "date", "years")
            stage = row_value(row, "stage", "description", "role", "institution")
        elif isinstance(row, (list, tuple)):
            period = str(row[0]).strip() if len(row) > 0 else ""
            stage = str(row[1]).strip() if len(row) > 1 else ""
        else:
            period = ""
            stage = str(row).strip()
        if period or stage:
            rows.append((period, stage))
    return rows


def article_rows(payload: dict[str, Any]) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for row in payload.get("articles", []):
        if isinstance(row, dict):
            year = row_value(row, "year", "publication_year")
            title = row_value(row, "title", "paper_title")
            journal = row_value(row, "journal", "venue", "source")
            role = row_value(row, "role", "scholar_role", "authorship_role")
        elif isinstance(row, (list, tuple)):
            values = [str(item).strip() for item in row]
            year = values[0] if len(values) > 0 else ""
            title = values[1] if len(values) > 1 else ""
            journal = values[2] if len(values) > 2 else ""
            role = values[3] if len(values) > 3 else ""
        else:
            continue
        if title:
            rows.append((year, title, journal, role))
    return rows


def escape_md_cell(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", "<br>")


def render_markdown(payload: dict[str, Any], path: Path, labels: dict[str, str]) -> None:
    title = str(require_payload(payload, "report_title")).strip()
    lines: list[str] = [f"# {title}", ""]

    lines.extend([f"## {labels['executive']}", ""])
    lines.extend(as_list(payload.get("executive_summary")) or [""])
    lines.append("")

    lines.extend([f"## {labels['identity']}", ""])
    lines.extend(as_list(payload.get("identity_and_anchors")) or [""])
    lines.append("")

    lines.extend([f"## {labels['timeline']}", ""])
    lines.append(f"| {labels['period']} | {labels['stage']} |")
    lines.append("|---|---|")
    for period, stage in timeline_rows(payload):
        lines.append(f"| {escape_md_cell(period)} | {escape_md_cell(stage)} |")
    if not timeline_rows(payload):
        lines.append("|  |  |")
    lines.append("")

    lines.extend([f"## {labels['articles']}", ""])
    lines.append(f"| {labels['year']} | {labels['title']} | {labels['journal']} | {labels['role']} |")
    lines.append("|---:|---|---|---|")
    for year, title_value, journal, role in article_rows(payload):
        lines.append(
            f"| {escape_md_cell(year)} | {escape_md_cell(title_value)} | "
            f"{escape_md_cell(journal)} | {escape_md_cell(role)} |"
        )
    if not article_rows(payload):
        lines.append("|  |  |  |  |")
    lines.append("")

    ensure_parent(path)
    path.write_text("\n".join(lines), encoding="utf-8", newline="\n")


def replace_paragraph_text(document: Document, placeholder: str, replacement: str) -> None:
    for paragraph in document.paragraphs:
        if placeholder not in paragraph.text:
            continue
        paragraph.text = paragraph.text.replace(placeholder, replacement)


def replace_heading_text(document: Document, old: str, new: str) -> None:
    if old == new:
        return
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old:
            paragraph.text = new


def set_cell_text(cell: Any, text: str) -> None:
    cell.text = str(text)


def clear_table_body(table: Any) -> None:
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)


def append_row_like(table: Any, values: list[str]) -> None:
    row = table.add_row()
    for idx, value in enumerate(values):
        set_cell_text(row.cells[idx], value)


def render_docx(payload: dict[str, Any], path: Path, labels: dict[str, str], template: Path) -> None:
    if not template.exists():
        raise SystemExit(f"Template not found: {template}")

    ensure_parent(path)
    shutil.copyfile(template, path)
    document = Document(path)

    title = str(require_payload(payload, "report_title")).strip()
    replace_paragraph_text(document, "{{REPORT_TITLE}}", title)
    replace_paragraph_text(document, "{{EXECUTIVE_SUMMARY}}", "\n".join(as_list(payload.get("executive_summary"))))
    replace_paragraph_text(document, "{{IDENTITY_AND_ANCHORS}}", "\n".join(as_list(payload.get("identity_and_anchors"))))

    replace_heading_text(document, "Executive summary", labels["executive"])
    replace_heading_text(document, "Scholar identity and anchors", labels["identity"])
    replace_heading_text(document, "Education and career timeline", labels["timeline"])
    replace_heading_text(
        document,
        "Qualifying first/co-first/corresponding/co-corresponding journal articles",
        labels["articles"],
    )

    if len(document.tables) < 2:
        raise SystemExit("Template must contain at least two tables")

    timeline_table = document.tables[0]
    article_table = document.tables[1]

    set_cell_text(timeline_table.rows[0].cells[0], labels["period"])
    set_cell_text(timeline_table.rows[0].cells[1], labels["stage"])
    clear_table_body(timeline_table)
    for period, stage in timeline_rows(payload):
        append_row_like(timeline_table, [period, stage])
    if not timeline_rows(payload):
        append_row_like(timeline_table, ["", ""])

    for idx, label_key in enumerate(("year", "title", "journal", "role")):
        set_cell_text(article_table.rows[0].cells[idx], labels[label_key])
    clear_table_body(article_table)
    for year, title_value, journal, role in article_rows(payload):
        append_row_like(article_table, [year, title_value, journal, role])
    if not article_rows(payload):
        append_row_like(article_table, ["", "", "", ""])

    document.save(path)
    Document(path)  # basic open validation


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Render authorship audit Markdown and DOCX reports.")
    parser.add_argument("payload", type=Path, help="UTF-8 JSON payload")
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE, help="DOCX template path")
    parser.add_argument("--md", type=Path, default=None, help="Override Markdown output path")
    parser.add_argument("--docx", type=Path, default=None, help="Override DOCX output path")
    args = parser.parse_args(argv)

    payload = json.loads(args.payload.read_text(encoding="utf-8-sig"))
    validate_articles(payload)
    language = normalize_language(payload.get("language"))
    labels = copy.deepcopy(LABELS[language])

    default_md_path, default_docx_path = default_output_paths(payload)
    md_path = args.md or Path(payload.get("output_md") or default_md_path)
    docx_path = args.docx or Path(payload.get("output_docx") or default_docx_path)

    render_markdown(payload, md_path, labels)
    render_docx(payload, docx_path, labels, args.template)

    print(f"Wrote Markdown: {md_path}")
    print(f"Wrote DOCX: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
